#!/usr/bin/env python3
"""Build a governed DMP trace from a protocol, DM log, template, and checklist."""

from __future__ import annotations

import argparse
import datetime as dt
import json
import re
from pathlib import Path
from typing import Any, Iterable


CHECKLIST_COLUMNS = [
    "序号",
    "规则文档章节/应用范围",
    "位置/主题",
    "非固定内容",
    "判断粒度",
    "统一判断/模板选择规则",
    "需要填写/替换的具体内容",
    "来源类型",
    "缺失时处理",
    "适用条件",
]

PROTECTED_TABLE_LIKE_SECTIONS = ["9", "15.2", "26.1", "27.1", "27.2", "27.3", "29"]

TEMPLATE_BY_DECISION = {
    "random": "DMP-随机系统.docx",
    "registry": "DMP-登记系统.docx",
    "none": "DMP-无随机无登记.docx",
}

DM_FIELD_ALIASES = {
    "临床试验方案名称": ["临床试验方案名称", "方案名称", "研究名称", "试验题目", "试验名称"],
    "方案名称": ["方案名称", "临床试验方案名称", "研究名称", "试验题目", "试验名称"],
    "方案编号": ["方案编号", "临床试验方案编号", "试验方案编号", "研究编号"],
    "申办者名称": ["申办者名称", "申办方名称", "申办者", "申办方", "申办单位"],
    "申办方名称": ["申办方名称", "申办者名称", "申办方", "申办者", "申办单位"],
    "数据管理单位名称": ["数据管理单位名称", "数据管理单位", "数据管理方名称", "数据管理方", "DM单位"],
    "数据管理单位": ["数据管理单位", "数据管理单位名称", "数据管理方名称", "数据管理方", "DM单位"],
    "临床监查方名称": ["临床监查方名称", "临床监察方名称", "临床监查方", "临床监察方", "监查方", "监察方", "CRO"],
    "临床监察方名称": ["临床监察方名称", "临床监查方名称", "临床监察方", "临床监查方", "监察方", "监查方", "CRO"],
    "版本号": ["DMP版本号", "版本号", "DMP版本"],
    "版本日期": ["DMP版本日期", "版本日期", "DMP日期"],
    "撰写者修订者": ["撰写者/修订者", "撰写者修订者", "撰写者", "修订者"],
    "是否使用随机系统": ["是否使用随机系统", "使用随机系统", "是否随机", "随机系统"],
    "是否使用登记系统": ["是否使用登记系统", "使用登记系统", "是否登记", "登记系统"],
    "统计分析单位名称": ["统计分析单位名称", "统计分析方名称", "统计分析单位", "统计分析方"],
    "统计分析方名称": ["统计分析方名称", "统计分析单位名称", "统计分析方", "统计分析单位"],
}

STRICT_IDENTIFIER_ITEMS = {
    "方案名称",
    "临床试验方案名称",
    "方案编号",
    "申办者名称",
    "申办方名称",
    "数据管理单位",
    "数据管理单位名称",
}

VERSION_RECORD_ALIASES = {
    "版本号": ["DMP版本号", "版本号", "DMP版本"],
    "版本日期": ["DMP版本日期", "版本日期", "DMP版本日期"],
    "撰写者/修订者": ["撰写者/修订者", "撰写者修订者", "撰写者", "修订者"],
    "版本修订内容": ["版本修订记录", "版本修订内容", "修订内容", "修订记录"],
}

SIGNER_FIELD_ALIASES = {
    "撰写人": ["撰写人", "撰写者", "撰写者/修订者", "撰写者修订者"],
    "数据管理单位审核人": ["数据管理单位审核人", "数据管理审核人", "DM审核人"],
    "申办者审核人": ["申办者审核人", "申办方审核人", "申办者复核人"],
    "CRO审核人": ["CRO审核人", "临床监查方审核人", "临床监察方审核人", "监查方审核人", "项目经理"],
    "统计分析单位审核人": ["统计分析单位审核人", "统计分析方审核人", "统计师审核人", "统计审核人"],
}

SIGNER_KEY_ORDER = [
    ("key 1", "数据管理单位审核人"),
    ("key 2", "申办者审核人"),
    ("key 3", "CRO审核人"),
    ("key 4", "统计分析单位审核人"),
]


def norm(value: Any) -> str:
    text = "" if value is None else str(value)
    return re.sub(r"[\s　：:，,。；;（）()、/\\_\-]+", "", text).lower()


def unique(values: Iterable[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        text = as_text(value)
        key = norm(text)
        if text and key not in seen:
            seen.add(key)
            result.append(text)
    return result


def aliases_for(item: str) -> list[str]:
    candidates = [item]
    item_norm = norm(item)
    for canonical, aliases in DM_FIELD_ALIASES.items():
        alias_norms = {norm(alias) for alias in [canonical, *aliases]}
        if item_norm in alias_norms:
            candidates.extend(aliases)
    return unique(candidates)


def as_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value).strip()


def load_xlsx_rows(path: Path) -> list[dict[str, Any]]:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    raw_rows = list(ws.iter_rows(values_only=True))
    if not raw_rows:
        return []
    headers = [as_text(v) for v in raw_rows[0]]
    rows = []
    for raw in raw_rows[1:]:
        row = {headers[i]: raw[i] if i < len(raw) else None for i in range(len(headers))}
        if any(as_text(v) for v in row.values()):
            rows.append(row)
    return rows


def _extract_xml_text_with_ins(element: Any) -> str:
    """Extract text from a docx XML element, including tracked insertions (w:ins)
    but excluding tracked deletions (w:del)."""
    from lxml import etree

    NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    text_parts: list[str] = []

    def _walk(el):
        tag = etree.QName(el).localname if isinstance(el.tag, str) else ""
        if tag == "del":
            return
        if tag == "ins":
            for child in el:
                _walk(child)
            return
        if tag == "t":
            txt = el.text or ""
            if txt:
                text_parts.append(txt)
            return
        if tag in ("rPr", "pPr", "pBdr", "tblPr", "tcPr", "trPr", "tblGrid"):
            return
        for child in el:
            _walk(child)

    _walk(element)
    return "".join(text_parts)


def read_docx_text(path: Path) -> tuple[str, list[str]]:
    from docx import Document

    doc = Document(path)
    lines: list[str] = []
    for para in doc.paragraphs:
        text = _extract_xml_text_with_ins(para._element).strip()
        if text:
            lines.append(text)
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            cells = [_extract_xml_text_with_ins(cell._tc).strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                lines.append(f"[表{table_index + 1} 行{row_index + 1}] " + " | ".join(cells))
    return "\n".join(lines), lines


def read_pdf_text(path: Path) -> tuple[str, list[str]]:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - optional dependency
        raise SystemExit("PDF input requires pypdf in the active Python environment.") from exc

    reader = PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        for line in (page.extract_text() or "").splitlines():
            line = line.strip()
            if line:
                lines.append(line)
    return "\n".join(lines), lines


def read_text_file(path: Path) -> tuple[str, list[str]]:
    text = path.read_text(encoding="utf-8")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return text, lines


def read_protocol(path: Path) -> tuple[str, list[str]]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx_text(path)
    if suffix == ".pdf":
        return read_pdf_text(path)
    if suffix in {".txt", ".md"}:
        return read_text_file(path)
    raise SystemExit(f"Unsupported protocol format: {path.suffix}")


def flatten_json(data: Any, prefix: str = "") -> dict[str, str]:
    flat: dict[str, str] = {}
    if isinstance(data, dict):
        for key, value in data.items():
            child_key = f"{prefix}.{key}" if prefix else str(key)
            if isinstance(value, (dict, list)):
                flat.update(flatten_json(value, child_key))
            else:
                flat[child_key] = as_text(value)
                flat[str(key)] = as_text(value)
    elif isinstance(data, list):
        for index, value in enumerate(data):
            flat.update(flatten_json(value, f"{prefix}[{index}]" if prefix else f"[{index}]"))
    else:
        flat[prefix or "value"] = as_text(data)
    return {key: value for key, value in flat.items() if value}


def read_dm_log(path: Path) -> tuple[dict[str, str], str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".json":
        data = json.loads(path.read_text(encoding="utf-8"))
        flat = flatten_json(data)
        return flat, json.dumps(data, ensure_ascii=False, indent=2), data
    if suffix == ".xlsx":
        rows = load_xlsx_rows(path)
        flat = flatten_json(rows)
        return flat, json.dumps(rows, ensure_ascii=False, indent=2), rows
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8")
        return {"DM日志全文": text}, text, text
    raise SystemExit(f"Unsupported DM log format: {path.suffix}")


def dm_lookup(item: str, dm_flat: dict[str, str], allow_fuzzy: bool = True) -> dict[str, Any]:
    aliases = aliases_for(item)
    for alias in aliases:
        if alias in dm_flat and dm_flat[alias]:
            return {
                "status": "filled",
                "value": dm_flat[alias],
                "evidence": [f"DM日志字段 `{alias}`: {dm_flat[alias]}"],
            }

    alias_norms = {norm(alias) for alias in aliases}
    for key, value in dm_flat.items():
        if norm(key) in alias_norms and value:
            return {"status": "filled", "value": value, "evidence": [f"DM日志字段 `{key}`: {value}"]}

    if not allow_fuzzy:
        return {"status": "missing", "value": None, "evidence": []}

    item_norm = norm(item)
    for key, value in dm_flat.items():
        key_norm = norm(key)
        if value and any(alias_norm in key_norm or key_norm in alias_norm for alias_norm in alias_norms | {item_norm}):
            return {"status": "uncertain", "value": value, "evidence": [f"DM日志近似字段 `{key}`: {value}"]}

    return {"status": "missing", "value": None, "evidence": []}


def mapping_lookup(mapping: dict[str, Any], aliases: Iterable[str]) -> str | None:
    alias_norms = {norm(alias) for alias in aliases}
    for key, value in mapping.items():
        if norm(key) in alias_norms and as_text(value):
            return as_text(value)
    return None


def has_version_fields(mapping: dict[str, Any]) -> bool:
    return any(mapping_lookup(mapping, aliases) for aliases in VERSION_RECORD_ALIASES.values())


def normalize_version_record(mapping: dict[str, Any]) -> dict[str, str]:
    record: dict[str, str] = {}
    for field, aliases in VERSION_RECORD_ALIASES.items():
        value = mapping_lookup(mapping, aliases)
        if value:
            record[field] = value
    return record


def _version_sort_key(record: dict[str, str]) -> tuple[str, str]:
    """Sort by version date then version number, both oldest first."""
    date_str = record.get("版本日期", "") or "0000-00-00"
    version_str = record.get("版本号", "") or ""
    date_normalized = re.sub(r"[-./]", "", date_str)
    return (date_normalized, version_str)


def extract_version_records(data: Any) -> list[dict[str, str]]:
    records: list[dict[str, str]] = []

    def visit(value: Any) -> None:
        if isinstance(value, list):
            for entry in value:
                if isinstance(entry, dict) and has_version_fields(entry):
                    records.append(normalize_version_record(entry))
                else:
                    visit(entry)
            return
        if isinstance(value, dict):
            if has_version_fields(value):
                records.append(normalize_version_record(value))
                return
            for child in value.values():
                visit(child)

    visit(data)
    deduped: list[dict[str, str]] = []
    seen: set[tuple[str, str, str, str]] = set()
    for record in records:
        key = (
            record.get("版本号", ""),
            record.get("版本日期", ""),
            record.get("撰写者/修订者", ""),
            record.get("版本修订内容", ""),
        )
        if key not in seen and any(key):
            seen.add(key)
            deduped.append(record)
    deduped.sort(key=_version_sort_key)
    return deduped


def signer_alias_norms() -> set[str]:
    aliases: list[str] = ["撰写人"]
    for role_key, role_aliases in SIGNER_FIELD_ALIASES.items():
        if role_key != "撰写人":
            aliases.extend(role_aliases)
    return {norm(alias) for alias in aliases}


def collect_signature_signers(mapping: dict[str, Any]) -> dict[str, str] | None:
    """Collect flat signer fields, with legacy nested signer-object fallback."""
    alias_norms = signer_alias_norms()
    found_signer_field = False
    signers: dict[str, str] = {}

    for key, raw_value in mapping.items():
        if isinstance(raw_value, (dict, list)):
            continue
        if norm(key) in alias_norms:
            found_signer_field = True
            signers[str(key).strip()] = as_text(raw_value)

    signer_block = mapping.get("签署页签署人")
    if isinstance(signer_block, dict):
        found_signer_field = True
        for key, raw_value in signer_block.items():
            text_key = str(key).strip()
            signers.setdefault(text_key, as_text(raw_value))

    if found_signer_field:
        return signers
    return None


def extract_signature_signers(data: Any) -> dict[str, str]:
    """Find current DM-log signature signer fields."""

    def visit(value: Any) -> dict[str, str] | None:
        if isinstance(value, dict):
            hit = collect_signature_signers(value)
            if hit is not None:
                return hit
            for child in value.values():
                hit = visit(child)
                if hit is not None:
                    return hit
        elif isinstance(value, list):
            for child in reversed(value):
                hit = visit(child)
                if hit is not None:
                    return hit
        return None

    return visit(data) or {}


def normalize_signers(raw_signers: dict[str, str]) -> dict[str, Any]:
    """Map DM-log signer names to the signature-page placeholder keys."""

    def lookup(role_key: str) -> str | None:
        aliases = SIGNER_FIELD_ALIASES.get(role_key, [role_key])
        alias_norms = {norm(alias) for alias in aliases}
        for key, value in raw_signers.items():
            if value and norm(key) in alias_norms:
                return as_text(value)
        return None

    reviewers: dict[str, dict[str, str]] = {}
    for placeholder_key, role_name in SIGNER_KEY_ORDER:
        reviewers[placeholder_key] = {
            "role": role_name,
            "name": lookup(role_name) or "",
        }

    return {
        "writer": lookup("撰写人") or "",
        "reviewers": reviewers,
        "raw": dict(raw_signers),
    }


def select_template_from_dm(dm_flat: dict[str, str], template_dir: Path) -> dict[str, Any]:
    random_result = dm_lookup("是否使用随机系统", dm_flat, allow_fuzzy=False)
    registry_result = dm_lookup("是否使用登记系统", dm_flat, allow_fuzzy=False)
    random_polarity = yes_no_polarity(random_result.get("value") or "")
    registry_polarity = yes_no_polarity(registry_result.get("value") or "")

    missing: list[str] = []
    if random_result["status"] not in {"filled", "uncertain"} or not random_polarity:
        missing.append("是否使用随机系统")
    if registry_result["status"] not in {"filled", "uncertain"} or not registry_polarity:
        missing.append("是否使用登记系统")
    if missing:
        details = "；".join(missing)
        raise SystemExit(
            "无法根据DM日志明确选择DMP模板，请先确认："
            f"{details}。模板选择必须先于DMP生成，且不能从方案静默推断。"
        )

    if random_polarity == "yes" and registry_polarity == "yes":
        raise SystemExit(
            "DM日志中「是否使用随机系统」和「是否使用登记系统」均为「是」，"
            "但项目中不会同时使用随机系统和登记系统。请确认DM日志中这两个字段的值是否正确，"
            "并提供修正后的值再重新生成DMP。"
        )

    if random_polarity == "yes":
        decision = "random"
    elif registry_polarity == "yes":
        decision = "registry"
    else:
        decision = "none"

    template_name = TEMPLATE_BY_DECISION[decision]
    template_path = template_dir / template_name
    if not template_path.exists():
        raise SystemExit(f"已根据DM日志选择模板 `{template_name}`，但文件不存在：{template_path}")

    return {
        "decision": decision,
        "template_path": template_path,
        "selected_template_name": template_name,
        "是否使用随机系统": random_result.get("value"),
        "是否使用登记系统": registry_result.get("value"),
        "evidence": random_result.get("evidence", []) + registry_result.get("evidence", []),
    }


def first_regex(patterns: Iterable[str], text: str) -> tuple[str | None, list[str]]:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.I | re.M)
        if match:
            value = match.group(1).strip()
            start = max(0, match.start() - 80)
            end = min(len(text), match.end() + 160)
            evidence = re.sub(r"\s+", " ", text[start:end]).strip()
            return value, [evidence]
    return None, []


def find_lines(lines: list[str], keywords: Iterable[str], limit: int = 3) -> list[str]:
    hits: list[str] = []
    for line in lines:
        if is_toc_like(line):
            continue
        if any(keyword in line for keyword in keywords):
            hits.append(line)
        if len(hits) >= limit:
            break
    return hits


def is_toc_like(line: str) -> bool:
    return bool(re.search(r"\t\d+$", line.strip())) or line.strip() in {"目录", "试验目的", "试验设计"}


def table_row_value(lines: list[str], labels: Iterable[str]) -> tuple[str | None, list[str]]:
    matches = table_row_values(lines, labels)
    if matches:
        _, value, evidence = matches[0]
        return value, evidence
    return None, []


def table_row_values(lines: list[str], labels: Iterable[str]) -> list[tuple[str, str, list[str]]]:
    label_norms = [norm(label) for label in labels]
    matches: list[tuple[str, str, list[str]]] = []
    for line in lines:
        if not line.startswith("[表"):
            continue
        body = re.sub(r"^\[表\d+\s+行\d+\]\s*", "", line, count=1)
        if " | " in body:
            label, value = body.split(" | ", 1)
            if semantic_label_matches(label, label_norms) and table_value_is_informative(value):
                matches.append((label.strip(), value.strip(), [line]))
                continue
        for label in labels:
            inline_match = re.match(rf"^{re.escape(label)}[：:]\s*(.+)$", body, flags=re.S)
            if inline_match:
                value = inline_match.group(1).strip()
                if table_value_is_informative(value):
                    matches.append((label, value, [line]))
                    break
    return matches


def semantic_label_matches(label: str, label_norms: Iterable[str]) -> bool:
    label_norm = norm(label)
    if not label_norm:
        return False
    for candidate in label_norms:
        if label_norm == candidate:
            return True
        if len(candidate) >= 4 and (candidate in label_norm or label_norm in candidate):
            return True
    return False


def table_value_is_informative(value: str) -> bool:
    text = value.strip()
    if not text:
        return False
    compact = re.sub(r"[\s|　:：,，。；;/-]+", "", text)
    compact = compact.replace("签名", "").replace("签章", "").replace("日期", "")
    compact = compact.replace("年月日", "").replace("年", "").replace("月", "").replace("日", "")
    if not compact:
        return False
    if compact.startswith("请输入"):
        return False
    return True


def effective_endpoint_from_summary(lines: list[str], endpoint_type: str) -> tuple[str | None, list[str]]:
    value, evidence = table_row_value(lines, ["有效性指标", "有效性终点", "疗效指标", "评价指标", "终点指标"])
    if not value:
        return None, []
    blocks = semantic_blocks(value)
    primary_start = find_first_index(blocks, is_primary_endpoint_block)
    other_start = find_first_index(blocks, is_other_endpoint_block, start=(primary_start + 1 if primary_start is not None else 0))
    if endpoint_type == "primary":
        if primary_start is not None:
            end = other_start if other_start is not None else len(blocks)
            return "\n".join(blocks[primary_start:end]).strip(), evidence
    if endpoint_type == "other":
        if other_start is not None:
            return "\n".join(blocks[other_start:]).strip(), evidence
    return None, evidence


def other_endpoints_from_summary(lines: list[str]) -> tuple[str | None, list[str]]:
    parts: list[str] = []
    evidence: list[str] = []

    value, value_evidence = effective_endpoint_from_summary(lines, "other")
    if value:
        parts.append(value)
        evidence.extend(value_evidence)

    additional_labels = [
        "次要终点",
        "次要研究终点",
        "次要有效性终点",
        "次要评价指标",
        "其他终点",
        "其他评价指标",
        "安全性指标",
        "安全性终点",
        "安全性评价指标",
        "安全性评价",
        "安全性观察指标",
        "探索性指标",
        "探索性终点",
        "探索性评价指标",
    ]
    for label, additional_value, additional_evidence in table_row_values(lines, additional_labels):
        formatted = f"{label}：\n{additional_value.strip()}"
        if not any(norm(formatted) == norm(existing) for existing in parts):
            parts.append(formatted)
            evidence.extend(additional_evidence)

    if parts:
        return "\n".join(part.strip() for part in parts if part.strip()), evidence
    return None, evidence


def semantic_blocks(value: str) -> list[str]:
    return [line.strip() for line in value.splitlines() if line.strip()]


def find_first_index(blocks: list[str], predicate, start: int = 0) -> int | None:
    for index in range(start, len(blocks)):
        if predicate(blocks[index]):
            return index
    return None


def is_primary_endpoint_block(block: str) -> bool:
    text = norm(block)
    return "主要" in text and any(term in text for term in ["终点", "指标", "评价"])


def is_other_endpoint_block(block: str) -> bool:
    text = norm(block)
    if not any(term in text for term in ["终点", "指标", "评价"]):
        return False
    return any(term in text for term in ["次要", "其他", "安全性", "探索"])


def analysis_population_value(lines: list[str]) -> tuple[str | None, list[str]]:
    evidence = []
    for line in lines:
        if re.search(r"(全分析集|Full Analysis Set|\bFAS\b)", line) and "SAMMPRIS" not in line:
            evidence.append(line)
        elif re.search(r"(符合方案集|Per Protocol Set|\bPPS\b)", line):
            evidence.append(line)
        elif re.search(r"(安全性数据集|安全性分析集|Safety Set|\bSS\b)", line) and "VISSIT" not in line:
            evidence.append(line)
        if len(evidence) >= 3:
            break
    if evidence:
        return "\n".join(evidence), evidence
    return None, []


def protocol_title(lines: list[str]) -> tuple[str | None, list[str]]:
    for line in lines[:40]:
        if len(line) < 12:
            continue
        if any(skip in line for skip in ["方案编号", "目录", "缩略语表", "临床试验流程图"]):
            continue
        if line.strip() in {"方案", "方 案"}:
            continue
        return line.strip(), [line.strip()]
    return None, []


def normalize_version(value: str | None) -> str | None:
    """Normalize version string to uppercase V prefix: v4.0→V4.0, 1.3→V1.3, V1.0→V1.0."""
    if not value:
        return value
    stripped = value.strip()
    if not stripped:
        return value
    if stripped[0].lower() == "v":
        return f"V{stripped[1:]}"
    if stripped[0].isdigit():
        return f"V{stripped}"
    return value


def filename_version(path: Path | None) -> tuple[str | None, list[str]]:
    if not path:
        return None, []
    match = re.search(r"(?i)(?:^|[^a-z0-9])v\s*(\d+(?:\.\d+)+)", path.stem)
    if not match:
        return None, []
    value = f"V{match.group(1)}"
    return value, [f"方案文件名 `{path.name}` 提供方案版本号线索：{value}"]


def filename_version_date(path: Path | None) -> tuple[str | None, list[str]]:
    if not path:
        return None, []
    # Try separated: 2024-01-11, 2024.01.11, 2024年01月11日
    match = re.search(r"(\d{4})[-_.年](\d{1,2})[-_.月](\d{1,2})", path.stem)
    if not match:
        # Try unseparated: 20240111 (8 consecutive digits)
        match = re.search(r"(?<!\d)(\d{4})(\d{2})(\d{2})(?!\d)", path.stem)
    if not match:
        return None, []
    y, m, d = int(match.group(1)), int(match.group(2)), int(match.group(3))
    if not (1 <= m <= 12 and 1 <= d <= 31):
        return None, []
    value = f"{y:04d}-{m:02d}-{d:02d}"
    return value, [f"方案文件名 `{path.name}` 提供方案版本日期线索：{value}"]


def normalize_date_value(value: str | None) -> str | None:
    if not value:
        return None
    match = re.search(r"(\d{4})[-./年](\d{1,2})[-./月](\d{1,2})", value)
    if not match:
        return value.strip()
    return f"{int(match.group(1)):04d}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"


def looks_like_partial_version(value: str | None, filename_value: str | None) -> bool:
    if not value or not filename_value:
        return False
    compact = value.rstrip(".")
    return filename_value.startswith(compact) and filename_value != value


def protocol_lookup(item: str, text: str, lines: list[str], protocol_path: Path | None = None) -> dict[str, Any]:
    if item in {"临床试验方案名称", "研究名称", "方案名称"}:
        value, evidence = table_row_value(lines, ["试验题目", "研究名称", "试验名称", "方案名称"])
        if value:
            return result_from_value(value, evidence, confident=True)
        value, evidence = protocol_title(lines)
        return result_from_value(value, evidence, confident=True)

    if item == "方案编号":
        value, evidence = first_regex([r"(?:临床试验)?方案编号[：:]\s*([^\s，,；;\n]+)"], text)
        return result_from_value(value, evidence, confident=True)

    if item == "方案版本号":
        value, evidence = first_regex(
            [
                r"(?:方案)?版本(?:号)?(?:[和及与/、]?(?:版本)?日期)?[：:]\s*(?:\|\s*)?([Vv]?\d+(?:\.\d+)+)",
                r"(?:方案)?版本(?:号)?(?:[和及与/、]?(?:版本)?日期)?[：:]\s*(?:\|\s*)?([Vv]?\d+(?:\.\d+)*)",
                r"Version\s*[:：]?\s*([Vv]?\d+(?:\.\d+)*)",
            ],
            text,
        )
        value = normalize_version(value)
        filename_value, filename_evidence = filename_version(protocol_path)
        filename_value = normalize_version(filename_value)
        if filename_value and (not value or looks_like_partial_version(value, filename_value)):
            return result_from_value(filename_value, filename_evidence + evidence, confident=True)
        if value and filename_value and norm(value) != norm(filename_value):
            return {
                "status": "conflict",
                "value": None,
                "evidence": evidence + filename_evidence,
            }
        return result_from_value(value, evidence, confident=bool(value))

    if item == "方案版本日期":
        value, evidence = first_regex(
            [
                # Combined version+date: 方案版本号和日期：V1.0/2025.11.24   or  版本号/版本日期：v4.0/2023-05-23
                r"(?:方案)?版本(?:号)?(?:和|及|与|\/|、)(?:版本)?日期[：:]\s*(?:V?\d+(?:\.\d+)*[/\\,;；，、]\s*)?([0-9]{4}[-./年][0-9]{1,2}[-./月][0-9]{1,2}日?)",
                # Standalone: 方案版本日期：2025.11.24
                r"(?:方案)?版本日期[：:]\s*([0-9]{4}[-./年][0-9]{1,2}[-./月][0-9]{1,2}日?)",
                # Table cell: V1.1，2024年9月2日 (date after version in same cell)
                r"(?:V\d+(?:\.\d+)*[/\\,;；，、]\s*)?([0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日)",
                # Bare date near 方案: 方案...日期：2025.11.24
                r"(?:方案)(?:.{0,8}?)(?:日期)[：:]\s*([0-9]{4}[-./年][0-9]{1,2}[-./月][0-9]{1,2}日?)",
            ],
            text,
        )
        filename_value, filename_evidence = filename_version_date(protocol_path)
        if value:
            normalized_value = normalize_date_value(value)
            result = result_from_value(normalized_value, evidence, confident=True)
            if filename_value and normalized_value != filename_value:
                result["evidence"].extend(filename_evidence)
                result["evidence"].append(
                    f"方案文件名日期 `{filename_value}` 与正文 `{normalized_value}` 不一致，以正文为准"
                )
            return result
        if filename_value:
            return {
                "status": "uncertain",
                "value": filename_value,
                "evidence": filename_evidence
                + ["注意：版本日期仅从文件名推断，未在方案正文中找到对应字段，请人工确认"],
            }
        return result_from_value(value, evidence, confident=bool(value))

    if item in {"申办者名称", "申办方名称"}:
        value, evidence = table_row_value(lines, ["申办者", "申办者名称", "申办方", "申办方名称", "申办单位"])
        if value:
            return result_from_value(value, evidence, confident=True)
        value, evidence = first_regex([r"申办者(?:名称|信息)?[：:]\s*([^\n]{2,80})"], text)
        if value:
            return result_from_value(value, evidence, confident=True)
        hits = find_lines(lines, ["申办者"], limit=5)
        return {"status": "uncertain" if hits else "missing", "value": None, "evidence": hits}

    if item in {"数据管理单位", "数据管理单位名称"}:
        value, evidence = table_row_value(lines, ["数据管理单位", "数据管理单位名称", "数据管理方", "数据管理方名称"])
        if value:
            return result_from_value(value, evidence, confident=True)
        value, evidence = first_regex([r"数据管理(?:单位|方)(?:名称)?[：:]\s*([^\n]{2,80})"], text)
        return result_from_value(value, evidence, confident=bool(value))

    if item in {"研究设计", "研究设计类型"}:
        value, evidence = table_row_value(lines, ["试验设计", "研究设计"])
        if value:
            return result_from_value(value, evidence, confident=True)
        hits = find_lines(lines, ["研究设计", "试验设计", "总体设计", "随机", "多中心", "单臂"], limit=6)
        title, title_evidence = protocol_title(lines)
        if title and any(word in title for word in ["随机", "多中心", "对照", "单臂", "盲", "优效", "非劣"]):
            return {"status": "filled", "value": title, "evidence": title_evidence}
        return value_from_hits(hits)

    if item == "研究目的":
        value, evidence = table_row_value(lines, ["试验目的", "研究目的"])
        if value:
            return result_from_value(value, evidence, confident=True)
        hits = find_lines(lines, ["研究目的", "试验目的", "主要目的", "次要目的"], limit=6)
        return value_from_hits(hits)

    if item == "样本量":
        value, evidence = table_row_value(lines, ["样本量"])
        if value:
            return result_from_value(value, evidence, confident=True)
        value, evidence = first_regex(
            [
                r"(?:样本量|计划入组|拟入组)[^。\n]{0,30}?([0-9０-９]+[^。\n；;]{0,40}(?:例|名|人|受试者))",
                r"([0-9０-９]+[^。\n；;]{0,20}(?:例|名|人|受试者))",
            ],
            text,
        )
        return result_from_value(value, evidence, confident=bool(value))

    if item == "主要有效性终点":
        value, evidence = effective_endpoint_from_summary(lines, "primary")
        if value:
            return result_from_value(value, evidence, confident=True)
        hits = find_lines(lines, ["主要有效性终点", "主要终点", "主要评价指标"], limit=8)
        return value_from_hits(hits)

    if item == "其他终点":
        value, evidence = other_endpoints_from_summary(lines)
        if value:
            return result_from_value(value, evidence, confident=True)
        hits = find_lines(lines, ["次要终点", "安全性终点", "探索性终点", "次要评价指标", "安全性评价指标"], limit=8)
        return value_from_hits(hits)

    if item == "统计分析人群":
        value, evidence = analysis_population_value(lines)
        if value:
            return result_from_value(value, evidence, confident=True)
        hits = find_lines(lines, ["全分析集", "符合方案集", "安全性分析集", "FAS", "PPS"], limit=8)
        return value_from_hits(hits)

    if item == "项目类型：药物 / 器械":
        if any(word in text for word in ["医疗器械", "器械临床试验", "支架系统", "医疗器械临床"]):
            return {"status": "filled", "value": "器械项目", "evidence": find_lines(lines, ["医疗器械", "器械", "支架系统"], limit=3)}
        if any(word in text for word in ["药物临床试验", "药品", "药物"]):
            return {"status": "filled", "value": "药物项目", "evidence": find_lines(lines, ["药物", "药品"], limit=3)}

    if item == "是否使用随机系统":
        hits = find_lines(lines, ["随机", "中央随机", "IWRS"], limit=6)
        if hits:
            return {"status": "filled", "value": "是", "evidence": hits}

    if item == "是否有阶段性分析/中期分析":
        hits = find_lines(lines, ["阶段性分析", "中期分析"], limit=6)
        if hits:
            return {"status": "filled", "value": "是", "evidence": hits}

    return {"status": "missing", "value": None, "evidence": []}


def result_from_value(value: str | None, evidence: list[str], confident: bool) -> dict[str, Any]:
    if value and confident:
        return {"status": "filled", "value": value, "evidence": evidence}
    if value:
        return {"status": "uncertain", "value": value, "evidence": evidence}
    return {"status": "missing", "value": None, "evidence": evidence}


def value_from_hits(hits: list[str]) -> dict[str, Any]:
    if not hits:
        return {"status": "missing", "value": None, "evidence": []}
    if len(hits) == 1:
        return {"status": "uncertain", "value": hits[0], "evidence": hits}
    return {"status": "uncertain", "value": "\n".join(hits), "evidence": hits}


def values_conflict(left: str | None, right: str | None) -> bool:
    if not left or not right:
        return False
    left_polarity = yes_no_polarity(left)
    right_polarity = yes_no_polarity(right)
    if left_polarity and right_polarity:
        return left_polarity != right_polarity
    return norm(left) != norm(right)


def yes_no_polarity(value: str) -> str | None:
    text = str(value).strip()
    negative_markers = ["否", "无", "不使用", "不涉及", "没有", "非随机", "不适用"]
    positive_markers = ["是", "有", "使用", "涉及", "随机"]
    if any(marker in text for marker in negative_markers):
        return "no"
    if any(marker in text for marker in positive_markers):
        return "yes"
    return None


def value_matches_expected(actual: str, expected: str) -> bool | None:
    expected_polarity = yes_no_polarity(expected)
    if expected_polarity:
        actual_polarity = yes_no_polarity(actual)
        if not actual_polarity:
            return None
        return actual_polarity == expected_polarity

    actual_norm = norm(actual)
    expected_norm = norm(expected)
    if not actual_norm or not expected_norm:
        return None
    return expected_norm in actual_norm or actual_norm in expected_norm


def evaluate_applicability_condition(condition: str, dm_flat: dict[str, str]) -> dict[str, Any]:
    condition = as_text(condition)
    if not condition:
        return {"status": "applicable", "evidence": []}

    parts = [part.strip() for part in re.split(r"[;；]", condition) if part.strip()]
    if not parts:
        return {"status": "applicable", "evidence": []}

    evidence: list[str] = []
    pending = False
    for part in parts:
        match = re.match(r"^(.+?)(?:=|＝|==)(.+)$", part)
        if not match:
            pending = True
            evidence.append(f"适用条件 `{part}` 无法解析，请使用 `字段名=期望值` 格式。")
            continue

        field = as_text(match.group(1))
        expected = as_text(match.group(2))
        result = dm_lookup(field, dm_flat, allow_fuzzy=False)
        if result["status"] not in {"filled", "uncertain"} or not result.get("value"):
            pending = True
            evidence.append(f"适用条件 `{part}` 无法判断：DM日志未提供 `{field}`。")
            continue

        actual = as_text(result.get("value"))
        matches = value_matches_expected(actual, expected)
        evidence.extend(result.get("evidence", []))
        if matches is None:
            pending = True
            evidence.append(f"适用条件 `{part}` 无法判断：`{field}` 当前值为 `{actual}`。")
        elif not matches:
            evidence.append(f"适用条件 `{part}` 不成立：`{field}` 当前值为 `{actual}`。")
            return {"status": "not_applicable", "evidence": evidence}
        else:
            evidence.append(f"适用条件 `{part}` 已满足。")

    if pending:
        return {"status": "condition_pending", "evidence": evidence}
    return {"status": "applicable", "evidence": evidence}


def make_question(row: dict[str, Any], status: str, evidence: list[str]) -> str:
    seq = as_text(row.get("序号"))
    section = as_text(row.get("规则文档章节/应用范围"))
    item = as_text(row.get("非固定内容"))
    required = as_text(row.get("需要填写/替换的具体内容")) or "请确认具体填写内容"
    source = as_text(row.get("来源类型")) or "用户确认"
    rule = as_text(row.get("统一判断/模板选择规则"))
    evidence_text = f" 当前线索：{' / '.join(evidence[:2])}" if evidence else ""
    return (
        f"[{seq}] {section} - {item}：请确认「{required}」。"
        f"来源应为：{source}。用途/规则：{rule}。当前状态：{status}。{evidence_text}"
    )


def choose_strict_identifier(
    protocol_result: dict[str, Any],
    dm_result: dict[str, Any],
    source_type: str,
) -> tuple[dict[str, Any], str]:
    protocol_has_value = protocol_result["status"] in {"filled", "uncertain"} and protocol_result.get("value")
    dm_has_value = dm_result["status"] in {"filled", "uncertain"} and dm_result.get("value")

    if protocol_has_value and dm_has_value and values_conflict(protocol_result.get("value"), dm_result.get("value")):
        return (
            {
                "status": "conflict",
                "value": None,
                "evidence": protocol_result.get("evidence", [])
                + [f"DM日志线索: {dm_result.get('value')}"]
                + dm_result.get("evidence", []),
            },
            "方案/DM日志",
        )
    if protocol_has_value:
        return protocol_result, "方案"
    if dm_has_value:
        return dm_result, "DM日志"
    if protocol_result.get("evidence"):
        return protocol_result, "方案"
    if dm_result.get("evidence"):
        return dm_result, "DM日志"
    if source_type == "暂不处理":
        return {"status": "missing", "value": None, "evidence": []}, "用户确认"
    return {"status": "missing", "value": None, "evidence": []}, source_type or "用户确认"


def compute_confidence(
    status: str,
    value: Any,
    evidence: list[str],
    source_used: str,
    source_type: str,
    extraction_method: str = "",
    applicable: bool = True,
) -> dict[str, Any]:
    """Compute confidence scores for a trace item.

    Returns dict with extraction_accuracy, completeness, hallucination_risk, overall_confidence.
    All scores 0-100. Low hallucination_risk = good (evidence-grounded).
    """
    # Defaults for non-applicable items
    if not applicable or status in {"not_applicable", "condition_pending", "not_processed"}:
        return {
            "extraction_accuracy": 0,
            "completeness": 0,
            "hallucination_risk": 0,
            "overall_confidence": 0,
            "extraction_method": extraction_method or "none",
            "scoring_note": "不适用",
        }

    # ---- extraction_accuracy ----
    accuracy = 50  # default
    if extraction_method == "dm_literal":
        accuracy = 95
    elif extraction_method == "dm_lookup":
        accuracy = 90
    elif extraction_method == "dm_fuzzy":
        accuracy = 70
    elif extraction_method in ("protocol_literal", "table_row"):
        accuracy = 90
    elif extraction_method == "protocol_regex":
        accuracy = 65
    elif extraction_method == "protocol_keyword":
        accuracy = 50
    elif extraction_method == "protocol_search":
        accuracy = 55
    elif extraction_method == "combined_fields":
        accuracy = 85
    elif extraction_method == "derived":
        accuracy = 70
    elif extraction_method == "user_confirm":
        accuracy = 40
    elif status == "filled" and not extraction_method:
        accuracy = 70
    elif status in {"missing", "uncertain", "conflict", "manual_confirm"}:
        accuracy = 10

    # ---- completeness ----
    completeness = 85  # default for filled items
    if status == "filled":
        if extraction_method in ("dm_literal", "dm_lookup", "protocol_literal", "table_row"):
            completeness = 90
        elif extraction_method in ("dm_fuzzy", "combined_fields"):
            completeness = 80
        elif extraction_method in ("protocol_regex", "protocol_keyword", "protocol_search"):
            completeness = 60
        elif extraction_method == "derived":
            completeness = 65
        elif extraction_method == "user_confirm":
            completeness = 50
        # Adjust based on evidence quality
        evidence_text = " ".join(evidence).lower() if evidence else ""
        if "截取" in evidence_text or "fragment" in evidence_text:
            completeness = min(completeness, 50)
        if "多行" in evidence_text or "multiple" in evidence_text:
            completeness = min(completeness, 70)
    elif status in {"uncertain", "conflict"}:
        completeness = 30
    elif status in {"missing", "manual_confirm"}:
        completeness = 10

    # ---- hallucination_risk (inverted: low = good) ----
    hallucination_risk = 20  # default: low risk for rule-based extraction
    if evidence:
        hallucination_risk = 5  # has evidence = very low risk
        if extraction_method in ("protocol_keyword", "protocol_search"):
            hallucination_risk = 25  # keyword-based evidence is less precise
    elif status == "filled" and not evidence:
        hallucination_risk = 60  # filled but no evidence = suspicious
    elif status in {"uncertain", "conflict", "manual_confirm"}:
        hallucination_risk = 40
    elif status == "missing":
        hallucination_risk = 70
    # source_used signals
    if "AI" in source_used or "推断" in source_used:
        hallucination_risk = max(hallucination_risk, 60)
    if "用户确认" in source_used:
        hallucination_risk = max(hallucination_risk, 50)

    # ---- overall_confidence ----
    # Weighted: 35% accuracy, 35% completeness, 30% (100 - hallucination_risk)
    anti_hallucination = max(0, 100 - hallucination_risk)
    overall = round(accuracy * 0.35 + completeness * 0.35 + anti_hallucination * 0.30)

    return {
        "extraction_accuracy": accuracy,
        "completeness": completeness,
        "hallucination_risk": hallucination_risk,
        "overall_confidence": overall,
        "extraction_method": extraction_method,
    }


def resolve_row(
    row: dict[str, Any],
    dm_flat: dict[str, str],
    protocol_text: str,
    protocol_lines: list[str],
    protocol_path: Path | None = None,
) -> dict[str, Any]:
    item = as_text(row.get("非固定内容"))
    source_type = as_text(row.get("来源类型"))
    granularity = as_text(row.get("判断粒度"))
    applicability_condition = as_text(row.get("适用条件"))

    applicability = evaluate_applicability_condition(applicability_condition, dm_flat)
    if applicability["status"] in {"not_applicable", "condition_pending"}:
        key = re.sub(r"[\s　：:，,。；;；()、/\\_-]+", "_", item).strip("_")
        confidence = compute_confidence(
            applicability["status"], None, applicability.get("evidence", []),
            "适用条件", source_type, applicable=False,
        )
        return {
            "seq": row.get("序号"),
            "key": key,
            "section": as_text(row.get("规则文档章节/应用范围")),
            "topic": as_text(row.get("位置/主题")),
            "item": item,
            "granularity": granularity,
            "rule": as_text(row.get("统一判断/模板选择规则")),
            "required_content": as_text(row.get("需要填写/替换的具体内容")),
            "source_type": source_type,
            "missing_handling": as_text(row.get("缺失时处理")),
            "applicability_condition": applicability_condition,
            "status": applicability["status"],
            "value": None,
            "source_used": "适用条件",
            "evidence": applicability.get("evidence", []),
            "question": None,
            "confidence": confidence,
        }

    extraction_method = ""  # tracked for confidence scoring

    if item == "页眉":
        required_text = as_text(row.get("需要填写/替换的具体内容"))
        if "DMP版本号" in required_text and "DMP版本日期" in required_text:
            version = dm_lookup("DMP版本号", dm_flat, allow_fuzzy=False)
            version_date = dm_lookup("DMP版本日期", dm_flat, allow_fuzzy=False)
            extraction_method = "combined_fields"
            if version["status"] == "filled" and version_date["status"] == "filled":
                selected = {
                    "status": "filled",
                    "value": f"{version['value']}，{version_date['value']}",
                    "evidence": version.get("evidence", []) + version_date.get("evidence", []),
                }
            else:
                selected = {
                    "status": "missing",
                    "value": None,
                    "evidence": version.get("evidence", []) + version_date.get("evidence", []),
                }
            source_used = "DM日志"
        elif "申办者" in required_text and "方案编号" in required_text:
            extraction_method = "combined_fields"
            sponsor = protocol_lookup("申办者名称", protocol_text, protocol_lines, protocol_path)
            protocol_no = protocol_lookup("方案编号", protocol_text, protocol_lines, protocol_path)
            if sponsor["status"] == "filled" and protocol_no["status"] == "filled":
                selected = {
                    "status": "filled",
                    "value": f"{sponsor['value']}，{protocol_no['value']}",
                    "evidence": sponsor.get("evidence", []) + protocol_no.get("evidence", []),
                }
            else:
                selected = {
                    "status": "missing",
                    "value": None,
                    "evidence": sponsor.get("evidence", []) + protocol_no.get("evidence", []),
                }
            source_used = "方案"
        else:
            selected = {"status": "not_processed", "value": None, "evidence": []}
            source_used = source_type or "未指定"
    elif item == "签署页配置":
        extraction_method = "dm_lookup"
        selected = dm_lookup("撰写者修订者", dm_flat, allow_fuzzy=False)
        source_used = "DM日志"
    else:
        dm_result = dm_lookup(item, dm_flat, allow_fuzzy=item not in STRICT_IDENTIFIER_ITEMS)
        protocol_result = protocol_lookup(item, protocol_text, protocol_lines, protocol_path)
        # Track extraction method based on lookup result quality
        if dm_result["status"] == "filled":
            evidence_str = " ".join(dm_result.get("evidence", []))
            if "近似" in evidence_str:
                extraction_method = "dm_fuzzy"
            else:
                extraction_method = "dm_lookup"
        elif protocol_result["status"] == "filled":
            extraction_method = "protocol_regex"
        else:
            extraction_method = "search"
        if item in STRICT_IDENTIFIER_ITEMS:
            selected, source_used = choose_strict_identifier(protocol_result, dm_result, source_type)
            extraction_method = "dm_literal" if source_used == "DM日志" else "protocol_literal"
        elif source_type == "DM日志":
            selected = dm_result
            source_used = "DM日志"
            if dm_result["status"] in {"filled", "uncertain"} and protocol_result["status"] == "filled":
                if values_conflict(dm_result["value"], protocol_result["value"]):
                    selected = {
                        "status": "conflict",
                        "value": None,
                        "evidence": dm_result["evidence"] + [f"方案线索: {protocol_result['value']}"] + protocol_result["evidence"],
                    }
                    source_used = "DM日志/方案"
        elif source_type == "方案":
            selected = protocol_result
            source_used = "方案"
            extraction_method = "protocol_regex"
            if protocol_result["status"] in {"filled", "uncertain"} and dm_result["status"] == "filled":
                if values_conflict(protocol_result["value"], dm_result["value"]):
                    selected = {
                        "status": "conflict",
                        "value": None,
                        "evidence": protocol_result["evidence"] + [f"DM日志线索: {dm_result['value']}"] + dm_result["evidence"],
                    }
                    source_used = "方案/DM日志"
        elif source_type == "暂不处理":
            extraction_method = "user_confirm"
            selected = {"status": "manual_confirm", "value": dm_result.get("value"), "evidence": dm_result.get("evidence", [])}
            source_used = "用户确认"
        elif not source_type and item == "数据录入和质疑模板":
            extraction_method = "derived"
            mode_result = dm_lookup("项目数据采集模式：EDC / PDC", dm_flat, allow_fuzzy=False)
            if mode_result["status"] == "filled" and mode_result.get("value"):
                selected = {
                    "status": "filled",
                    "value": f"{mode_result['value']}模板",
                    "evidence": mode_result.get("evidence", []),
                }
            else:
                selected = {"status": "missing", "value": None, "evidence": mode_result.get("evidence", [])}
            source_used = "DM日志"
        elif not source_type and item in {"AI可自动判断的项目特征", "方案与DM日志信息冲突", "方案和DM日志均未提供的信息"}:
            selected = {"status": "qc_rule", "value": None, "evidence": []}
            source_used = "QC规则"
        else:
            selected = {"status": "not_processed", "value": None, "evidence": []}
            source_used = source_type or "未指定"

    status = selected["status"]
    question = None
    missing_rule = as_text(row.get("缺失时处理"))
    if status in {"missing", "uncertain", "conflict", "manual_confirm", "not_processed"} and missing_rule != "NA":
        question = make_question(row, status, selected.get("evidence", []))

    key = re.sub(r"[\s　：:，,。；;；()、/\\_-]+", "_", item).strip("_")
    confidence = compute_confidence(
        status, selected.get("value"), selected.get("evidence", []),
        source_used, source_type, extraction_method, applicable=True,
    )
    return {
        "seq": row.get("序号"),
        "key": key,
        "section": as_text(row.get("规则文档章节/应用范围")),
        "topic": as_text(row.get("位置/主题")),
        "item": item,
        "granularity": granularity,
        "rule": as_text(row.get("统一判断/模板选择规则")),
        "required_content": as_text(row.get("需要填写/替换的具体内容")),
        "source_type": source_type,
        "missing_handling": missing_rule,
        "applicability_condition": applicability_condition,
        "status": status,
        "value": selected.get("value"),
        "source_used": source_used,
        "evidence": selected.get("evidence", []),
        "question": question,
        "confidence": confidence,
    }


def write_questions(path: Path, trace_items: list[dict[str, Any]]) -> None:
    questions = [item for item in trace_items if item.get("question")]
    lines = ["# DMP待确认问题", ""]
    if not questions:
        lines.append("无待确认问题。")
    else:
        current_section = None
        for item in questions:
            section = item.get("section") or "未指定章节"
            if section != current_section:
                lines.extend(["", f"## {section}"])
                current_section = section
            lines.append(f"- {item['question']}")
    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a traceable DMP non-fixed-content map.")
    parser.add_argument("--protocol", required=True, type=Path)
    parser.add_argument("--dm-log", required=True, type=Path)
    parser.add_argument("--template", type=Path)
    parser.add_argument("--template-dir", type=Path)
    parser.add_argument("--checklist", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--questions", type=Path)
    parser.add_argument("--protocol-dump", type=Path)
    args = parser.parse_args()

    dm_flat, dm_raw, dm_data = read_dm_log(args.dm_log)
    version_records = extract_version_records(dm_data)
    signature_signers = normalize_signers(extract_signature_signers(dm_data))

    dm_entry_count = 1
    if isinstance(dm_data, list):
        dm_entry_count = len(dm_data)
    elif isinstance(dm_data, dict) and any(
        isinstance(v, list) for v in dm_data.values()
    ):
        dm_entry_count = max(
            len(v) for v in dm_data.values() if isinstance(v, list)
        )

    if args.template_dir:
        template_selection = select_template_from_dm(dm_flat, args.template_dir)
        template_path = template_selection["template_path"]
    elif args.template:
        template_path = args.template
        template_selection = {
            "decision": "explicit",
            "template_path": template_path,
            "selected_template_name": template_path.name,
            "是否使用随机系统": dm_lookup("是否使用随机系统", dm_flat, allow_fuzzy=False).get("value"),
            "是否使用登记系统": dm_lookup("是否使用登记系统", dm_flat, allow_fuzzy=False).get("value"),
            "evidence": ["使用命令行显式模板参数。"],
        }
    else:
        raise SystemExit("请提供 --template-dir 以便根据DM日志选择模板，或显式提供 --template。")

    checklist_rows = load_xlsx_rows(args.checklist)
    protocol_text, protocol_lines = read_protocol(args.protocol)

    if args.protocol_dump:
        args.protocol_dump.write_text(protocol_text, encoding="utf-8")

    items = [resolve_row(row, dm_flat, protocol_text, protocol_lines, args.protocol) for row in checklist_rows]
    trace = {
        "metadata": {
            "created_at": dt.datetime.now().isoformat(timespec="seconds"),
            "protocol": str(args.protocol),
            "dm_log": str(args.dm_log),
            "template": str(template_path),
            "selected_template_name": template_path.name,
            "template_selection": {
                key: value
                for key, value in template_selection.items()
                if key not in {"template_path"}
            },
            "checklist": str(args.checklist),
            "checklist_rows": len(checklist_rows),
            "protected_table_like_sections": PROTECTED_TABLE_LIKE_SECTIONS,
            "version_records": version_records,
            "dm_entry_count": dm_entry_count,
            "signature_signers": signature_signers,
        },
        "items": items,
        "dm_log_keys": sorted(dm_flat.keys()),
    }

    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(trace, ensure_ascii=False, indent=2), encoding="utf-8")

    if args.questions:
        args.questions.parent.mkdir(parents=True, exist_ok=True)
        write_questions(args.questions, items)

    summary: dict[str, int] = {}
    for item in items:
        summary[item["status"]] = summary.get(item["status"], 0) + 1
    print(json.dumps(summary, ensure_ascii=False, sort_keys=True))


if __name__ == "__main__":
    main()
