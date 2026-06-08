#!/usr/bin/env python3
"""Semantic review of high-risk trace fields before applying to the DMP template.

Two modes:
  prepare  – extract review context from trace + protocol for LLM audit
  apply    – apply LLM corrections back to the trace JSON
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

# Fields whose extracted values are prone to semantic errors
HIGH_RISK_FIELDS = [
    "样本量",
    "研究设计",
    "主要有效性终点",
    "其他终点",
    "统计分析人群",
]

# How many chars of protocol context to capture around evidence
CONTEXT_WINDOW = 1000


def norm(value: str) -> str:
    return re.sub(r"[\s　：:,。；;()()、/\\_\-]+", "", str(value)).lower()


def _extract_xml_text_with_ins(element: Any) -> str:
    """Extract text from a docx XML element, including tracked insertions (w:ins)
    but excluding tracked deletions (w:del)."""
    from lxml import etree

    text_parts: list[str] = []

    def _walk(el):
        tag = etree.QName(el).localname if isinstance(el.tag, str) else ""
        if tag == "del":
            return
        if tag == "ins":
            for child in el:
                _walk(child)
            return
        if tag == "t":
            txt = el.text or ""
            if txt:
                text_parts.append(txt)
            return
        if tag in ("rPr", "pPr", "pBdr", "tblPr", "tcPr", "trPr", "tblGrid"):
            return
        for child in el:
            _walk(child)

    _walk(element)
    return "".join(text_parts)


def read_protocol_text(protocol_path: Path) -> str:
    suffix = protocol_path.suffix.lower()
    if suffix == ".docx":
        from docx import Document
        doc = Document(protocol_path)
        parts: list[str] = []
        for para in doc.paragraphs:
            text = _extract_xml_text_with_ins(para._element).strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [_extract_xml_text_with_ins(cell._tc).strip() for cell in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(protocol_path))
        return "\n".join(
            line.strip()
            for page in reader.pages
            for line in (page.extract_text() or "").splitlines()
            if line.strip()
        )
    return protocol_path.read_text(encoding="utf-8")


def find_evidence_span(full_text: str, evidence_snippet: str, window: int = CONTEXT_WINDOW) -> str:
    """Locate evidence snippet in full text and return surrounding context."""
    # Try exact match first, then fuzzy
    idx = full_text.find(evidence_snippet)
    if idx == -1:
        # Try matching the first 60 meaningful chars
        clean = evidence_snippet.strip()[:80]
        idx = full_text.find(clean)
    if idx == -1:
        return evidence_snippet  # fallback
    start = max(0, idx - window // 2)
    end = min(len(full_text), idx + len(evidence_snippet) + window // 2)
    return full_text[start:end]


def cmd_prepare(trace_path: Path, protocol_path: Path, protocol_text_path: Path | None, out_path: Path) -> None:
    trace = json.loads(trace_path.read_text(encoding="utf-8"))
    # Prefer pre-dumped text file (fast) over re-parsing docx (slow).
    if protocol_text_path:
        protocol_text = protocol_text_path.read_text(encoding="utf-8")
    else:
        protocol_text = read_protocol_text(protocol_path)
    items = trace.get("items", [])

    review_items: list[dict] = []
    for item in items:
        if item.get("item") not in HIGH_RISK_FIELDS:
            continue
        if item.get("status") not in {"filled", "uncertain"}:
            continue

        evidence_list = item.get("evidence", [])
        evidence_text = evidence_list[0] if evidence_list else ""
        context = find_evidence_span(protocol_text, evidence_text) if evidence_text else ""

        review_items.append({
            "key": item["key"],
            "item": item["item"],
            "section": item.get("section", ""),
            "current_value": item.get("value"),
            "current_status": item.get("status"),
            "evidence_snippet": evidence_text[:300],
            "protocol_context": context,
            # LLM fills these:
            "corrected_value": None,
            "correction_reason": "",
            "review_decision": "",  # "accept" | "correct" | "flag"
        })

    out = {
        "metadata": trace.get("metadata", {}),
        "protocol_path": str(protocol_path),
        "review_items": review_items,
    }
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    stats = ", ".join(f"{ri['item']}({ri['current_value'][:30] if ri['current_value'] else 'None'})" for ri in review_items)
    print(f"Prepared {len(review_items)} review items: {stats}")


def cmd_apply(review_path: Path, trace_path: Path, out_path: Path) -> None:
    review = json.loads(review_path.read_text(encoding="utf-8"))
    trace = json.loads(trace_path.read_text(encoding="utf-8"))
    items = trace.get("items", [])

    corrections: dict[str, dict] = {}
    for ri in review.get("review_items", []):
        if ri.get("review_decision") == "correct" and ri.get("corrected_value"):
            corrections[ri["key"]] = ri

    if not corrections:
        print("No corrections found in review – trace unchanged.")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(json.dumps(trace, ensure_ascii=False, indent=2), encoding="utf-8")
        return

    for item in items:
        key = item.get("key")
        if key in corrections:
            corr = corrections[key]
            old_value = item.get("value")
            new_value = corr["corrected_value"]
            item["value"] = new_value
            item["status"] = "filled"
            item["source_used"] = item.get("source_used", "方案") + " + LLM语义审核"
            old_evidence = item.get("evidence", [])
            item["evidence"] = old_evidence + [
                f"LLM语义审核修正: {corr.get('correction_reason', '语义审核纠正')}",
                f"修正前值: {old_value}",
            ]
            item["question"] = None  # clear stale question from pre-correction status
            print(f"CORRECTED [{key}]: {old_value} → {new_value}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(trace, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Applied {len(corrections)} corrections to trace.")


def main() -> None:
    parser = argparse.ArgumentParser(description="Semantic review of high-risk DMP trace fields")
    parser.add_argument("--mode", required=True, choices=["prepare", "apply"])
    parser.add_argument("--trace", required=True, type=Path, help="dmp_trace.json path")
    parser.add_argument("--protocol", type=Path, help="protocol docx/pdf (for semantic review context)")
    parser.add_argument("--protocol-text", type=Path, help="pre-dumped protocol text file (fast path)")
    parser.add_argument("--review", type=Path, help="review JSON path")
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()

    if args.mode == "prepare":
        if not args.protocol and not args.protocol_text:
            raise SystemExit("--protocol or --protocol-text is required for prepare mode")
        review_out = args.out
        cmd_prepare(args.trace, args.protocol, args.protocol_text, review_out)
    else:
        review_in = args.review or args.out
        cmd_apply(review_in, args.trace, args.out)


if __name__ == "__main__":
    main()
