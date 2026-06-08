#!/usr/bin/env python3
"""Combined semantic review + few-shot format constraint for DMP trace fields.

Merges semantic_review.py and fewshot_format.py into a single prepare→review→apply
cycle, so the LLM reviews all fields in one pass instead of two separate rounds.

Two modes:
  prepare  – extract combined review context (semantic + few-shot) from trace
  apply    – apply both semantic corrections and few-shot reformats to trace

When a field needs both semantic review and few-shot formatting, corrections are
applied first, then reformats run on the corrected values — matching the original
two-pass order.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# High-risk fields prone to regex extraction errors (from semantic_review.py)
# ---------------------------------------------------------------------------
HIGH_RISK_FIELDS = [
    "样本量",
    "研究设计",
    "主要有效性终点",
    "其他终点",
    "统计分析人群",
]

CONTEXT_WINDOW = 1000


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def norm(value: str) -> str:
    return re.sub(r"[\s　：:,。；;()()、/\\_\-]+", "", str(value)).lower()


# ---------------------------------------------------------------------------
# Protocol reading (from semantic_review.py)
# ---------------------------------------------------------------------------

def _extract_xml_text_with_ins(element: Any) -> str:
    """Extract text from a docx XML element, including tracked insertions (w:ins)
    but excluding tracked deletions (w:del)."""
    from lxml import etree

    text_parts: list[str] = []

    def _walk(el):
        tag = etree.QName(el).localname if isinstance(el.tag, str) else ""
        if tag == "del":
            return
        if tag == "ins":
            for child in el:
                _walk(child)
            return
        if tag == "t":
            txt = el.text or ""
            if txt:
                text_parts.append(txt)
            return
        if tag in ("rPr", "pPr", "pBdr", "tblPr", "tcPr", "trPr", "tblGrid"):
            return
        for child in el:
            _walk(child)

    _walk(element)
    return "".join(text_parts)


def read_protocol_text(protocol_path: Path) -> str:
    suffix = protocol_path.suffix.lower()
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(protocol_path))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = _extract_xml_text_with_ins(para._element).strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [_extract_xml_text_with_ins(cell._tc).strip() for cell in row.cells]
                cells = [c for c in cells if c]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(protocol_path))
        return "\n".join(
            line.strip()
            for page in reader.pages
            for line in (page.extract_text() or "").splitlines()
            if line.strip()
        )
    return protocol_path.read_text(encoding="utf-8")


def find_evidence_span(
    full_text: str, evidence_snippet: str, window: int = CONTEXT_WINDOW
) -> str:
    """Locate evidence snippet in full text and return surrounding context."""
    idx = full_text.find(evidence_snippet)
    if idx == -1:
        clean = evidence_snippet.strip()[:80]
        idx = full_text.find(clean)
    if idx == -1:
        return evidence_snippet
    start = max(0, idx - window // 2)
    end = min(len(full_text), idx + len(evidence_snippet) + window // 2)
    return full_text[start:end]


# ---------------------------------------------------------------------------
# Few-shot parsing (from fewshot_format.py)
# ---------------------------------------------------------------------------

def parse_fewshot(
    fewshot_path: Path,
) -> tuple[dict[str, list[str]], dict[str, str]]:
    """Parse fewshot.md into ({field_name: [examples]}, {field_name: output_rule_prompt})."""
    text = fewshot_path.read_text(encoding="utf-8")
    sections: dict[str, list[str]] = {}
    prompts: dict[str, str] = {}

    PROMPT_HEADER_RE = re.compile(r"^(.+?)的输出规范(?:（[^）]*）)?[：:]")

    lines = text.split("\n")
    header_positions: list[tuple[int, str, bool]] = []  # (line_idx, field_name, is_prompt)
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        prompt_match = PROMPT_HEADER_RE.match(stripped)
        if prompt_match:
            header_positions.append((i, prompt_match.group(1).strip(), True))
            continue
        field_match = re.match(r"^(.+?)(?:示例)[：:]$", stripped)
        if not field_match:
            field_match = re.match(r"^(.{2,8})[：:]$", stripped)
        if field_match:
            field_name = field_match.group(1).strip()
            if re.match(r"^(?:示例|例)\s*\d+$", field_name):
                continue
            if field_name and "注意" not in field_name and "输出规范" not in field_name:
                header_positions.append((i, field_name, False))

    for idx, (start_line, field_name, is_prompt) in enumerate(header_positions):
        header_line = lines[start_line].strip()
        if is_prompt:
            colon_pos = header_line.find("：")
            if colon_pos == -1:
                colon_pos = header_line.find(":")
            header_body = header_line[colon_pos + 1:].strip() if colon_pos >= 0 else ""
            content_start = start_line + 1
        else:
            header_body = ""
            content_start = start_line + 1

        if idx + 1 < len(header_positions):
            content_end = header_positions[idx + 1][0]
        else:
            content_end = len(lines)

        content_parts: list[str] = []
        if header_body:
            content_parts.append(header_body)
        for j in range(content_start, content_end):
            content_parts.append(lines[j])

        content = "\n".join(content_parts).strip()
        if not content:
            continue

        if is_prompt:
            prompts[field_name] = content
        else:
            examples: list[str] = []
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                match = re.match(r"(?:示例\s*\d+|例\s*\d+)[：:]\s*(.+)", line)
                if match:
                    examples.append(match.group(1).strip())
                elif not line.startswith("示例") and not line.startswith("例"):
                    examples.append(line)
            if examples:
                sections[field_name] = examples

    return sections, prompts


# ---------------------------------------------------------------------------
# Prepare mode
# ---------------------------------------------------------------------------

def cmd_prepare(
    trace_path: Path,
    protocol_path: Path | None,
    protocol_text_path: Path | None,
    fewshot_path: Path | None,
    out_path: Path,
) -> None:
    trace = json.loads(trace_path.read_text(encoding="utf-8"))
    items = trace.get("items", [])

    # --- Determine semantic-review targets ---
    semantic_targets: set[str] = set()
    needs_protocol = bool(protocol_path or protocol_text_path)
    if needs_protocol:
        for item in items:
            if (
                item.get("item") in HIGH_RISK_FIELDS
                and item.get("status") in {"filled", "uncertain"}
            ):
                semantic_targets.add(item["item"])

    # --- Determine few-shot targets ---
    fewshot_examples: dict[str, list[str]] = {}
    fewshot_prompts: dict[str, str] = {}
    if fewshot_path:
        fewshot_examples, fewshot_prompts = parse_fewshot(fewshot_path)

    fewshot_targets: set[str] = set()
    if fewshot_path:
        # example-based targets
        for field_name in fewshot_examples:
            for item in items:
                if item.get("item") == field_name and item.get("status") in {
                    "filled",
                    "uncertain",
                }:
                    fewshot_targets.add(field_name)
                    break
        # prompt-based targets (output-rule prompts without examples)
        for field_name in fewshot_prompts:
            if field_name in fewshot_targets:
                continue
            for item in items:
                if item.get("item") == field_name and item.get("status") in {
                    "filled",
                    "uncertain",
                }:
                    fewshot_targets.add(field_name)
                    break

    # --- Read protocol text once if semantic review is needed ---
    # Prefer pre-dumped text file (fast) over re-parsing docx (slow).
    protocol_text = ""
    protocol_source = None
    if semantic_targets:
        if protocol_text_path:
            protocol_text = protocol_text_path.read_text(encoding="utf-8")
            protocol_source = str(protocol_text_path)
        elif protocol_path:
            protocol_text = read_protocol_text(protocol_path)
            protocol_source = str(protocol_path)

    # --- Build combined review items ---
    all_targets = semantic_targets | fewshot_targets
    review_items: list[dict] = []
    seen: set[str] = set()

    for item in items:
        item_name = item.get("item", "")
        if item_name not in all_targets:
            continue
        if item_name in seen:
            continue
        seen.add(item_name)

        needs_semantic = item_name in semantic_targets
        needs_fewshot = item_name in fewshot_targets

        # Semantic review context
        evidence_snippet = ""
        protocol_context = ""
        if needs_semantic:
            evidence_list = item.get("evidence", [])
            evidence_text = evidence_list[0] if evidence_list else ""
            evidence_snippet = evidence_text[:300] if evidence_text else ""
            protocol_context = (
                find_evidence_span(protocol_text, evidence_text)
                if evidence_text
                else ""
            )

        # Few-shot context
        examples = fewshot_examples.get(item_name, [])
        prompt = fewshot_prompts.get(item_name, "")

        # Current confidence scores from trace
        current_confidence = item.get("confidence", {})

        review_items.append(
            {
                "key": item["key"],
                "item": item_name,
                "section": item.get("section", ""),
                "current_value": item.get("value"),
                "current_status": item.get("status"),
                # Semantic review fields
                "needs_semantic_review": needs_semantic,
                "evidence_snippet": evidence_snippet,
                "protocol_context": protocol_context,
                "corrected_value": None,
                "correction_reason": "",
                "review_decision": "",  # "accept" | "correct" | "flag" | "" (n/a)
                # Few-shot format fields
                "needs_fewshot_format": needs_fewshot,
                "fewshot_examples": examples,
                "field_prompt": prompt,
                "formatted_value": None,
                "format_reason": "",
                "format_decision": "",  # "accept" | "reformat" | "flag" | "" (n/a)
                # Confidence adjustment — CLI sets these during review
                "current_confidence": {
                    "extraction_accuracy": current_confidence.get("extraction_accuracy", 0),
                    "completeness": current_confidence.get("completeness", 0),
                    "hallucination_risk": current_confidence.get("hallucination_risk", 0),
                    "overall_confidence": current_confidence.get("overall_confidence", 0),
                },
                "confidence_adjustment": {
                    "extraction_accuracy_delta": 0,
                    "completeness_delta": 0,
                    "hallucination_risk_delta": 0,
                    "adjustment_reason": "",
                },
            }
        )

    out = {
        "metadata": trace.get("metadata", {}),
        "protocol_path": protocol_source,
        "fewshot_path": str(fewshot_path) if fewshot_path else None,
        "semantic_fields": sorted(semantic_targets),
        "fewshot_fields": sorted(fewshot_targets),
        "review_items": review_items,
    }
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")

    if review_items:
        stats = ", ".join(
            f"{ri['item']}({str(ri['current_value'])[:30] if ri['current_value'] else 'None'})"
            for ri in review_items
        )
        print(
            f"Prepared {len(review_items)} review items "
            f"(semantic: {len(semantic_targets)}, few-shot: {len(fewshot_targets)}): "
            f"{stats}"
        )
    else:
        print("No review items to prepare.")


# ---------------------------------------------------------------------------
# Apply mode
# ---------------------------------------------------------------------------

def cmd_apply(review_path: Path, trace_path: Path, out_path: Path) -> None:
    review = json.loads(review_path.read_text(encoding="utf-8"))
    trace = json.loads(trace_path.read_text(encoding="utf-8"))
    items = trace.get("items", [])

    # Collect semantic corrections
    corrections: dict[str, dict] = {}
    for ri in review.get("review_items", []):
        if ri.get("review_decision") == "correct" and ri.get("corrected_value"):
            corrections[ri["key"]] = ri

    # Collect few-shot reformats
    reformats: dict[str, dict] = {}
    for ri in review.get("review_items", []):
        if ri.get("format_decision") == "reformat" and ri.get("formatted_value"):
            reformats[ri["key"]] = ri

    correction_count = 0
    reformat_count = 0

    # Step 1: Apply semantic corrections first
    for item in items:
        key = item.get("key")
        if key in corrections:
            corr = corrections[key]
            old_value = item.get("value")
            new_value = corr["corrected_value"]
            item["value"] = new_value
            item["status"] = "filled"
            item["source_used"] = (
                item.get("source_used", "方案") + " + LLM语义审核"
            )
            old_evidence = item.get("evidence", [])
            item["evidence"] = old_evidence + [
                f"LLM语义审核修正: {corr.get('correction_reason', '语义审核纠正')}",
                f"修正前值: {old_value}",
            ]
            item["question"] = None
            print(
                f"CORRECTED [{key}]: {str(old_value)[:60]} → {str(new_value)[:60]}"
            )
            correction_count += 1

    # Step 2: Apply few-shot reformats on (possibly corrected) values
    for item in items:
        key = item.get("key")
        if key in reformats:
            ref = reformats[key]
            old_value = item.get("value")
            new_value = ref["formatted_value"]
            item["value"] = new_value
            item["status"] = "filled"
            item["source_used"] = (
                item.get("source_used", "方案") + " + few-shot格式化"
            )
            old_evidence = item.get("evidence", [])
            item["evidence"] = old_evidence + [
                f"few-shot格式化: {ref.get('format_reason', '按fewshot示例格式约束输出')}",
                f"格式化前值: {old_value}",
            ]
            item["question"] = None
            print(
                f"REFORMATTED [{key}]: {str(old_value)[:60]}... → {str(new_value)[:60]}..."
            )
            reformat_count += 1

    # Step 3: Apply confidence adjustments
    confidence_adjustments: dict[str, dict] = {}
    for ri in review.get("review_items", []):
        adj = ri.get("confidence_adjustment", {})
        if adj and any([
            adj.get("extraction_accuracy_delta", 0) != 0,
            adj.get("completeness_delta", 0) != 0,
            adj.get("hallucination_risk_delta", 0) != 0,
        ]):
            confidence_adjustments[ri["key"]] = adj

    adjustment_count = 0
    for item in items:
        key = item.get("key")
        if key in confidence_adjustments:
            adj = confidence_adjustments[key]
            conf = item.get("confidence", {})
            if not conf:
                continue

            old_acc = conf.get("extraction_accuracy", 0)
            old_comp = conf.get("completeness", 0)
            old_hallu = conf.get("hallucination_risk", 0)

            new_acc = max(0, min(100, old_acc + adj.get("extraction_accuracy_delta", 0)))
            new_comp = max(0, min(100, old_comp + adj.get("completeness_delta", 0)))
            new_hallu = max(0, min(100, old_hallu + adj.get("hallucination_risk_delta", 0)))

            anti_hallu = max(0, 100 - new_hallu)
            new_overall = round(new_acc * 0.35 + new_comp * 0.35 + anti_hallu * 0.30)

            conf["extraction_accuracy"] = new_acc
            conf["completeness"] = new_comp
            conf["hallucination_risk"] = new_hallu
            conf["overall_confidence"] = new_overall
            reason = adj.get("adjustment_reason", "")
            if reason:
                conf["scoring_note"] = (conf.get("scoring_note", "") + "; " + reason).strip("; ")

            item["confidence"] = conf
            print(
                f"CONFIDENCE [{key}]: acc {old_acc}→{new_acc}, comp {old_comp}→{new_comp}, "
                f"hallu {old_hallu}→{new_hallu}, overall→{new_overall}"
            )
            adjustment_count += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(
        json.dumps(trace, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(
        f"Applied {correction_count} corrections and {reformat_count} reformats to trace."
    )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Combined semantic review + few-shot format for DMP trace fields"
    )
    parser.add_argument(
        "--mode", required=True, choices=["prepare", "apply"]
    )
    parser.add_argument("--trace", required=True, type=Path, help="dmp_trace.json path")
    parser.add_argument(
        "--protocol", type=Path, help="protocol docx/pdf (for semantic review context)"
    )
    parser.add_argument(
        "--protocol-text",
        type=Path,
        help="pre-dumped protocol text file (fast path, avoids re-parsing docx)",
    )
    parser.add_argument(
        "--fewshot", type=Path, help="fewshot.md path (for format examples)"
    )
    parser.add_argument(
        "--review",
        type=Path,
        help="review JSON path (for apply, defaults to --out)",
    )
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()

    if args.mode == "prepare":
        cmd_prepare(args.trace, args.protocol, args.protocol_text, args.fewshot, args.out)
    else:
        review_in = args.review or args.out
        cmd_apply(review_in, args.trace, args.out)


if __name__ == "__main__":
    main()
